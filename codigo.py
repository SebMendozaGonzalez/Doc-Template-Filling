import pandas as pd
from docx import Document
from typing import Dict


def create_docs(df: pd.DataFrame) -> None:
    """This function revceives a dataframe and, for each row in the dataframe, fills up a word document template
        with the row's info and saves it in the "docs" folder

    Args:
        df (pandas dataframe): 
    """
    
    for i in range(len(df)):
        doc_copy = Document("./document_template.docx")
        
        law_firm: str = df.iloc[i,0]
        law_firm_street: str = df.iloc[i,1]
        law_firm_city: str = df.iloc[i,2]
        patient_name: str = df.iloc[i, 3]
        date_of_loss: str = df.iloc[i, 4]
        date_of_service: str = df.iloc[i, 5]
        balance: str = df.iloc[i, 6]
        
        dic1: Dict[str,str] = {
            "{Insert Law Firm}": law_firm,
            "{Insert Law Firm Street}": law_firm_street,
            "{Insert Law Firm City, State, Zip}": law_firm_city
        }
        
        for holder, value in dic1.items():
            for para in doc_copy.paragraphs:
                if para.text.startswith(holder):
                    para.text = f"{value}"
                    
        
        dic2: Dict[str,str] = {
            "Patient/Plaintiff: ": patient_name,
            "Date of Loss: ": date_of_loss,
            "Date of Service(s): ": date_of_service,
            "Balance Due: $": balance
        }
        
        for holder, value in dic2.items():
            for para in doc_copy.paragraphs:
                if para.text.startswith(holder):
                    para.text = f"{holder}{value}"
                elif para.text.startswith("This letter shall serve as notice that AD Pharmacy holds a balance with your client, "):
                    para.text = f"This letter shall serve as notice that AD Pharmacy holds a balance with your client, {dic2['Patient/Plaintiff: ']}."
                    
                    
        doc_copy.save(f"./documentos/modified{i}.docx")


if __name__ == "__main__":
    df: pd.DataFrame = pd.read_excel("./AD Pharm.xlsx")

    create_docs(df)